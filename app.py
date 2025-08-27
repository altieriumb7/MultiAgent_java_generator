
import os, sys, io, json, zipfile, tempfile, subprocess, shutil
from pathlib import Path
from typing import Optional, Dict, Any, List
import streamlit as st

# ----------------------------------------------------------------------------------
# Questo file deve stare nella root del repo, accanto a 'src'.
# Aggiungo la root al PYTHONPATH per importare src.app.*
# ----------------------------------------------------------------------------------
HERE = Path(__file__).resolve().parent
if str(HERE) not in sys.path:
    sys.path.insert(0, str(HERE))

# Importa i moduli esistenti del progetto
from src.app.tools.parser import parse_rfa_to_spec, parse_rfa_to_spec_llm
from src.app.tools.patterns import mine_project_style
from src.app.tools.codegen import render_artifacts
from src.app.tools.testgen import render_tests
from src.app.tools.fs import write_artifacts
from src.app.llm import have_openai, get_env_model, chat_complete

st.set_page_config(page_title="RFA → Java/Spring Code (progetto esistente)", layout="wide")
st.title("📦 RFA → Java/Spring Code Generator (sul tuo progetto esistente)")
st.caption("Carica un RFA/spec e genera DTO/Service/Controller/Client/Mapper + JUnit usando il tuo codice attuale.")

# ---------------------------- Helpers ----------------------------
def _read_upload_to_text(upload) -> str:
    name = upload.name.lower()
    data = upload.read()

    # .md / .txt
    if name.endswith(('.md', '.txt')):
        return data.decode('utf-8', errors='ignore')

    # .docx
    if name.endswith('.docx'):
        try:
            from docx import Document  # python-docx
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            tmp.write(data); tmp.close()
            doc = Document(tmp.name)
            os.unlink(tmp.name)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return data.decode('utf-8', errors='ignore')

    # .pdf
    if name.endswith('.pdf'):
        try:
            from pypdf import PdfReader
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            tmp.write(data); tmp.close()
            reader = PdfReader(tmp.name)
            pages = []
            for pg in reader.pages:
                try: pages.append(pg.extract_text() or "")
                except Exception: pass
            os.unlink(tmp.name)
            return "\n".join(pages).strip() or data.decode('utf-8', errors='ignore')
        except Exception:
            return data.decode('utf-8', errors='ignore')

    # .doc (Word 97-2003) — tentativi in ordine: Word (COM), LibreOffice, textract
    if name.endswith('.doc'):
        txt = ""

        # A) Microsoft Word via COM (Windows + pywin32)
        try:
            import win32com.client as win32  # pip install pywin32
            tmpdir = tempfile.mkdtemp()
            src = Path(tmpdir) / "input.doc"
            dst = Path(tmpdir) / "input.docx"
            src.write_bytes(data)

            word = win32.gencache.EnsureDispatch('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(str(src))
            doc.SaveAs(str(dst), FileFormat=16)  # 16 = wdFormatXMLDocument (docx)
            doc.Close(False)
            word.Quit()

            from docx import Document
            d = Document(str(dst))
            txt = "\n".join(p.text for p in d.paragraphs)
            shutil.rmtree(tmpdir, ignore_errors=True)
            if txt.strip():
                return txt
        except Exception:
            pass

        # B) LibreOffice (soffice) se installato
        try:
            tmpdir = tempfile.mkdtemp()
            src = Path(tmpdir) / "input.doc"
            dst = Path(tmpdir) / "input.docx"
            src.write_bytes(data)
            candidates = [
                os.environ.get("SOFFICE_PATH"),
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                "soffice",
            ]
            for exe in candidates:
                if not exe:
                    continue
                try:
                    cp = subprocess.run([exe, "--headless", "--convert-to", "docx", "--outdir", tmpdir, str(src)],
                                        stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, timeout=60)
                    if cp.returncode == 0 and dst.exists():
                        from docx import Document
                        d = Document(str(dst))
                        txt = "\n".join(p.text for p in d.paragraphs)
                        break
                except Exception:
                    pass
            shutil.rmtree(tmpdir, ignore_errors=True)
            if txt.strip():
                return txt
        except Exception:
            pass

        # C) textract (richiede dipendenze esterne su Windows)
        try:
            import textract  # pip install textract
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".doc")
            tmp.write(data); tmp.close()
            txt = textract.process(tmp.name).decode("utf-8", "ignore")
            os.unlink(tmp.name)
            if txt.strip():
                return txt
        except Exception:
            pass

        # Fallback
        return data.decode('utf-8', errors='ignore')

    # Fallback per altri formati
    return data.decode('utf-8', errors='ignore')


def _zip_dir(folder: Path) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for p in folder.rglob('*'):
            if p.is_file():
                zf.write(p, p.relative_to(folder))
    buf.seek(0)
    return buf.read()

# ---------------------------- Sidebar ----------------------------
with st.sidebar:
    st.subheader("Percorsi progetto")
    default_root = str(HERE)
    project_root = st.text_input("Root progetto Maven (cartella con pom.xml)", value=default_root)
    default_out = str(Path(project_root) / "generated") if project_root else default_root
    out_dir_in = st.text_input("Cartella output (predefinito: <root>/generated)", value=default_out)

    st.markdown("---")
    st.subheader("LLM")
    llm_mode = st.radio("Uso LLM", ["Auto (se API key)", "Sempre ON", "Sempre OFF"], index=0)
    api_key = st.text_input("OPENAI_API_KEY", value=os.environ.get("OPENAI_API_KEY",""), type="password")
    base_url = st.text_input("OPENAI_BASE_URL (opzionale)", value=os.environ.get("OPENAI_BASE_URL",""))
    doc_model = st.text_input("Doc model", value=os.environ.get("DOC_MODEL","gpt-4o-mini"), key="doc_model")

    if api_key:
        os.environ["OPENAI_API_KEY"] = api_key
    if base_url:
        os.environ["OPENAI_BASE_URL"] = base_url
    elif "OPENAI_BASE_URL" in os.environ:
        del os.environ["OPENAI_BASE_URL"]

    st.caption(f"LLM attivo: {'Sì' if have_openai() else 'No'} (in base alla presenza di OPENAI_API_KEY)")

    st.markdown("---")
    st.subheader("Override stile (opzionale)")
    force_lombok = st.selectbox("Lombok", ["auto", "force_on", "force_off"], index=0)
    force_feign = st.selectbox("Feign",  ["auto", "force_on", "force_off"], index=0)
    force_records = st.selectbox("Records",["auto", "force_on", "force_off"], index=0)

# ---------------------------- Generazione ----------------------------
st.header("Genera")
upload = st.file_uploader("Carica RFA/spec (.md, .txt, .docx, .pdf, .doc)", type=["md","txt","docx","pdf","doc"])

if st.button("Genera", type="primary", disabled=not upload):
    # Leggi doc
    rfa_text = _read_upload_to_text(upload)
    if not rfa_text.strip():
        st.error("Il file caricato non contiene testo leggibile.")
        st.stop()

    # Percorsi
    pr = Path(project_root).resolve() if project_root else None
    out_dir = Path(out_dir_in).resolve() if out_dir_in else (pr or HERE)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Determina uso LLM
    if llm_mode == "Sempre OFF":
        use_llm_flag = False
    elif llm_mode == "Sempre ON":
        use_llm_flag = True
    else:
        use_llm_flag = have_openai()

    # Parsing
    st.write("Parsing RFA…")
    if use_llm_flag:
        if not have_openai():
            st.error("Hai richiesto LLM ma OPENAI_API_KEY non è impostata.")
            st.stop()
        model = st.session_state.get("doc_model") or get_env_model("DOC_MODEL", "gpt-4o-mini")
        spec = parse_rfa_to_spec_llm(rfa_text, model=model)
        st.success(f"Parsing con LLM ({model})")
    else:
        spec = parse_rfa_to_spec(rfa_text)
        st.info("Parsing con parser euristico (senza LLM)")

    # Miner stile
    st.write("Analisi stile progetto…")
    style = mine_project_style(pr) if pr else {}
    if force_lombok != "auto":
        style["use_lombok"] = (force_lombok == "force_on")
    if force_feign != "auto":
        style["use_feign"] = (force_feign == "force_on")
    if force_records != "auto":
        style["use_records"] = (force_records == "force_on")

    # Code + test
    st.write("Generazione codice e test…")
    artifacts = render_artifacts(spec, style)
    tests = render_tests(spec, style)
    all_files = {**artifacts, **tests}
    write_artifacts(out_dir, all_files)

    st.success("Fatto!")
    st.json({"style": style, "files_written": len(all_files), "llm_used": bool(use_llm_flag)})
    try:
        data = _zip_dir(out_dir)
        st.download_button("Scarica output (ZIP)", data=data, file_name="generated_artifacts.zip")
    except Exception as e:
        st.warning(f"ZIP non riuscito: {e}")

# ---------------------------- Build (facoltativo) ----------------------------
st.header("Build (opzionale)")
with st.expander("Esegui Maven", expanded=False):
    goals = st.text_input("Maven goals", value="clean package -DskipTests")
    if st.button("Esegui Maven"):
        if not project_root or not Path(project_root).exists():
            st.error("Root progetto non valida.")
        else:
            try:
                cp = subprocess.run(["mvn"] + goals.split(), cwd=str(Path(project_root)), stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
                ok = cp.returncode == 0
                st.code(cp.stdout[-10000:], language="bash")
                st.success("Build OK") if ok else st.error("Build fallita")
            except FileNotFoundError:
                st.error("Maven (mvn) non trovato nel PATH.")

# ---------------------------- Fixer (facoltativo, richiede LLM) ----------------------------
st.header("Fixer Chat (opzionale, richiede LLM)")
if not os.environ.get("OPENAI_API_KEY"):
    st.info("Inserisci OPENAI_API_KEY nella sidebar per attivare il fixer.")
else:
    msg = st.text_area("Incolla un errore o fai una domanda; il fixer proporrà patch (file completi).", height=140)
    if st.button("Proponi & Applica Patch", disabled=not msg):
        system = "Sei un senior Java/Spring. Rispondi con JSON che contiene patch di file interi per risolvere il problema."
        schema = """Restituisci SOLO questo oggetto JSON:
{"patches":[{"path":"src/main/java/...","content":"<contenuto completo del file>"}],"note":"breve spiegazione"}"""
        user = f"Project root: {project_root}\nUser:\n{msg}\n{schema}"
        try:
            raw = chat_complete(
                model=get_env_model("CODE_MODEL", get_env_model("DOC_MODEL","gpt-4o-mini")),
                system=system,
                user=user,
                temperature=0.1,
                response_format={"type":"json_object"}
            )
            data = json.loads(raw)
            patched = []
            for p in data.get("patches", []):
                rel = Path(p.get("path","")).as_posix().strip("/")
                if not rel: continue
                dst = Path(project_root) / rel
                dst.parent.mkdir(parents=True, exist_ok=True)
                dst.write_text(p.get("content",""), encoding="utf-8")
                patched.append(rel)
            st.success(f"Applicata/e {len(patched)} patch.")
            if patched: st.json(patched)
            if data.get("note"): st.caption(data["note"])
        except Exception as e:
            st.error(f"Fixer fallito: {e}")
